from __future__ import annotations

import importlib
import json
from pathlib import Path
import zipfile
import zlib

from pypdf import PdfReader


def dump_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def write_png(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    raw_scanline = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw_scanline)

    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
        return len(data).to_bytes(4, "big") + chunk_type + data + crc.to_bytes(4, "big")

    png_bytes = b"".join(
        [
            b"\x89PNG\r\n\x1a\n",
            chunk(b"IHDR", b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00"),
            chunk(b"IDAT", compressed),
            chunk(b"IEND", b""),
        ]
    )
    path.write_bytes(png_bytes)


def write_docx(path: Path, text: str) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(text)
    document.sections[0].footer.paragraphs[0].text = f"{text} footer"
    document.save(path)


def make_paper_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "build" / "review_manuscript.md",
        """---
title: "Test Medical Manuscript"
bibliography: ../references.bib
link-citations: true
---

# Abstract

This is a manuscript citation [@ref1].

# Main Figures

## Figure 1. Main figure

Caption.

![](../figures/F1_main.png)

## Supplementary Figure S1. Supplementary figure

Caption.

![](../figures/FS1_supp.png)

# Main Tables

| Characteristic | Value |
| --- | --- |
| Age | 52 |
""",
    )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_png(paper_root / "figures" / "FS1_supp.png")
    write_text(paper_root / "figures" / "FS1_supp.pdf", "%PDF-1.4\n%supp figure\n")
    write_text(paper_root / "tables" / "T1_summary.csv", "Characteristic,Value\nAge,52\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "source_markdown": "paper/build/review_manuscript.md",
            "output_pdf": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "export_paths": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                },
                {
                    "figure_id": "FS1",
                    "paper_role": "supplementary",
                    "title": "Supplementary figure",
                    "export_paths": ["paper/figures/FS1_supp.pdf", "paper/figures/FS1_supp.png"],
                },
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "title": "Summary table",
                    "asset_paths": ["paper/tables/T1_summary.csv", "paper/tables/T1_summary.md"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "bundle_inputs": {
                "compiled_markdown_path": "paper/build/review_manuscript.md",
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
                {"path": "paper/figures/F1_main.pdf", "kind": "figure_export", "status": "present"},
                {"path": "paper/figures/FS1_supp.pdf", "kind": "supplementary_figure_export", "status": "present"},
                {"path": "paper/tables/T1_summary.csv", "kind": "table_asset", "status": "present"},
            ],
        },
    )

    return paper_root


def make_current_draft_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "draft.md",
        """# Draft

## Title

Current Draft Title

## Abstract

### Aims

Draft abstract aim with evidence [@ref1].

### Methods

Draft abstract methods.

### Results

Draft abstract results.

### Conclusions

Draft abstract conclusions.

## Introduction

Intro paragraph with citation [@ref1].

## Methods

Methods paragraph.

## Results

Results paragraph.

## Discussion

Discussion paragraph.

## Conclusion

Conclusion paragraph.
""",
        )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A. and Author, B.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "pdf_path": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "planned_exports": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "path": "paper/tables/T1_summary.md",
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/draft.md",
            "compile_report_path": "paper/build/compile_report.json",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
            ],
        },
    )
    return paper_root


def make_manuscript_shaped_draft_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "draft.md",
        """# Manuscript-Shaped Draft Title

## Abstract

### Background

Draft abstract background with evidence [@ref1].

### Methods

Draft abstract methods.

### Results

Draft abstract results.

### Conclusions

Draft abstract conclusions.

## Introduction

Intro paragraph with citation [@ref1].

## Materials and Methods

Study methods paragraph.

## Results

Results paragraph.

## Discussion

Discussion paragraph.
""",
    )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A. and Author, B.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "pdf_path": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "planned_exports": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "path": "paper/tables/T1_summary.md",
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/draft.md",
            "compile_report_path": "paper/build/compile_report.json",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
            ],
        },
    )
    return paper_root


def test_create_submission_minimal_package_creates_output_directory_and_copies_pdf(tmp_path: Path) -> None:
    try:
        module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    except ModuleNotFoundError:
        module = None

    assert module is not None
    paper_root = make_paper_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    assert submission_root.exists()
    assert (submission_root / "paper.pdf").exists()
    assert manifest["output_root"] == "paper/submission_minimal"


def test_create_submission_minimal_package_writes_manifest_and_docx_path(tmp_path: Path) -> None:
    try:
        module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    except ModuleNotFoundError:
        module = None

    assert module is not None
    paper_root = make_paper_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    manifest_path = submission_root / "submission_manifest.json"
    docx_path = submission_root / "manuscript.docx"

    assert manifest_path.exists()
    assert docx_path.exists()

    manifest_payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest_payload["publication_profile"] == "general_medical_journal"
    assert manifest_payload["citation_style"] == "AMA"
    assert manifest_payload["output_root"] == "paper/submission_minimal"
    assert manifest_payload["manuscript"]["pdf_path"] == "paper/submission_minimal/paper.pdf"
    assert manifest_payload["manuscript"]["docx_path"] == "paper/submission_minimal/manuscript.docx"
    assert manifest_payload["naming_map"]["figures"] == {
        "F1": "Figure1",
        "FS1": "SupplementaryFigureS1",
    }
    assert manifest_payload["naming_map"]["tables"] == {
        "T1": "Table1",
    }
    assert manifest_payload["figures"][0]["source_paths"] == [
        "paper/figures/F1_main.pdf",
        "paper/figures/F1_main.png",
    ]
    assert manifest_payload["tables"][0]["source_paths"] == [
        "paper/tables/T1_summary.csv",
        "paper/tables/T1_summary.md",
    ]
    assert manifest_payload == manifest


def test_create_submission_minimal_package_defaults_to_ama_citation_style(tmp_path: Path) -> None:
    try:
        module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    except ModuleNotFoundError:
        module = None

    assert module is not None
    paper_root = make_paper_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    assert manifest["citation_style"] == "AMA"


def test_create_submission_minimal_package_copies_figures_and_tables(tmp_path: Path) -> None:
    try:
        module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    except ModuleNotFoundError:
        module = None

    assert module is not None
    paper_root = make_paper_workspace(tmp_path)

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    expected_paths = [
        submission_root / "figures" / "Figure1.pdf",
        submission_root / "figures" / "Figure1.png",
        submission_root / "figures" / "SupplementaryFigureS1.pdf",
        submission_root / "figures" / "SupplementaryFigureS1.png",
        submission_root / "tables" / "Table1.csv",
        submission_root / "tables" / "Table1.md",
    ]

    for path in expected_paths:
        assert path.exists(), path


def test_create_submission_minimal_package_general_profile_writes_figure_legends_and_tables(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    dump_json(
        paper_root / "figure_semantics_manifest.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "Figure1",
                    "story_role": "overall_performance_and_clinical_utility",
                    "research_question": "Does the main figure support the paper-facing interpretation?",
                    "direct_message": "The primary display item supports the manuscript-facing clinical message.",
                    "clinical_implication": "The figure can be read as a reviewer-facing legend rather than a slide-style caption.",
                    "interpretation_boundary": "The figure legend does not establish a treatment recommendation by itself.",
                    "panel_messages": [
                        {"panel_id": "A", "message": "Panel A summarizes the main paper-facing interpretation."}
                    ],
                    "legend_glossary": [
                        {
                            "term": "treat all",
                            "explanation": "Assumes every patient is managed as high risk at the chosen threshold."
                        }
                    ],
                    "threshold_semantics": "Thresholds are illustrative operating points rather than mandated cut-offs.",
                    "stratification_basis": "Displayed groups follow the prespecified manuscript presentation.",
                    "recommendation_boundary": "Clinical decisions should not rely on this figure alone.",
                    "renderer_contract": {
                        "figure_semantics": "evidence",
                        "renderer_family": "python",
                        "selection_rationale": "The legend is derived from an audited paper-facing figure.",
                        "fallback_on_failure": False,
                        "failure_action": "block_and_fix_environment"
                    }
                }
            ],
        },
    )

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_markdown = (paper_root / "submission_minimal" / "manuscript_submission.md").read_text(encoding="utf-8")
    assert "# Figure Legends" in submission_markdown
    assert "## Figure 1. Main figure" in submission_markdown
    assert "Caption." in submission_markdown
    assert "The primary display item supports the manuscript-facing clinical message." in submission_markdown
    assert "Panel A summarizes the main paper-facing interpretation." in submission_markdown
    assert (
        "Abbreviations: treat all, Assumes every patient is managed as high risk at the chosen threshold."
        in submission_markdown
    )
    assert "# Tables" in submission_markdown
    assert "## Table 1" in submission_markdown
    assert "| Characteristic | Value |" in submission_markdown


def test_create_submission_minimal_package_general_profile_embeds_figures_into_docx_and_pdf(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    submission_markdown = (submission_root / "manuscript_submission.md").read_text(encoding="utf-8")
    assert "# Figures" in submission_markdown
    assert "## Figure 1. Main figure" in submission_markdown
    assert "![](../figures/F1_main.png)" in submission_markdown
    assert "Caption." in submission_markdown
    assert manifest["manuscript"]["pdf_path"] == "paper/submission_minimal/paper.pdf"

    with zipfile.ZipFile(submission_root / "manuscript.docx") as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    assert any(name.startswith("word/media/") for name in names)
    assert "<w:drawing" in document_xml

    pdf_reader = PdfReader(str(submission_root / "paper.pdf"))
    assert sum(len(page.images) for page in pdf_reader.pages) >= 1


def test_create_submission_minimal_package_accepts_current_bundle_contract_shape(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "pdf_path": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/build/review_manuscript.md",
            "compile_report_path": "paper/build/compile_report.json",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
            ],
        },
    )

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    assert (submission_root / "manuscript.docx").exists()
    assert (submission_root / "paper.pdf").exists()
    assert manifest["manuscript"]["source_markdown_path"] == "paper/submission_minimal/manuscript_submission.md"
    assert manifest["manuscript"]["pdf_path"] == "paper/submission_minimal/paper.pdf"


def test_create_submission_minimal_package_prefers_compiled_markdown_over_draft_path(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    write_text(
        paper_root / "draft.md",
        """# Draft

## Title

Wrong draft title

## Abstract

Wrong draft abstract.
""",
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/draft.md",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
        },
    )

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_markdown = (paper_root / "submission_minimal" / "manuscript_submission.md").read_text(encoding="utf-8")
    assert 'title: "Test Medical Manuscript"' in submission_markdown
    assert 'bibliography: ../references.bib' in submission_markdown
    assert "Wrong draft title" not in submission_markdown


def test_create_submission_minimal_package_accepts_current_figure_and_table_catalog_shape(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "figure1",
                    "role": "paper_main",
                    "planned_exports": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "table1",
                    "path": "paper/tables/T1_summary.md",
                }
            ],
        },
    )

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    assert (submission_root / "figures" / "figure1.pdf").exists()
    assert (submission_root / "figures" / "figure1.png").exists()
    assert (submission_root / "tables" / "table1.md").exists()


def test_create_submission_minimal_package_skips_missing_planned_table_entries(tmp_path: Path) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)

    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "table1",
                    "status": "rendered_and_cleaned",
                    "path": "paper/tables/T1_summary.md",
                },
                {
                    "table_id": "table3",
                    "status": "planned_from_trusted_reports",
                    "path": "paper/tables/T3_missing.md",
                },
            ],
        },
    )

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    assert (submission_root / "tables" / "table1.md").exists()
    assert not (submission_root / "tables" / "table3.md").exists()


def test_create_submission_minimal_package_syncs_study_delivery_when_context_is_available(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    called: dict[str, object] = {}

    def fake_can_sync(*, paper_root: Path) -> bool:
        called["can_sync_paper_root"] = paper_root
        return True

    def fake_sync(
        *,
        paper_root: Path,
        stage: str,
        publication_profile: str = "general_medical_journal",
    ) -> dict:
        called["sync_paper_root"] = paper_root
        called["sync_stage"] = stage
        called["sync_publication_profile"] = publication_profile
        return {"stage": stage, "publication_profile": publication_profile}

    monkeypatch.setattr(module.study_delivery_sync, "can_sync_study_delivery", fake_can_sync)
    monkeypatch.setattr(module.study_delivery_sync, "sync_study_delivery", fake_sync)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    assert called["can_sync_paper_root"] == paper_root
    assert called["sync_paper_root"] == paper_root
    assert called["sync_stage"] == "submission_minimal"
    assert called["sync_publication_profile"] == "general_medical_journal"
    assert manifest["delivery_sync"] == {
        "stage": "submission_minimal",
        "publication_profile": "general_medical_journal",
    }
    assert manifest["readme_path"] == "paper/submission_minimal/README.md"
    readme_text = (paper_root / "submission_minimal" / "README.md").read_text(encoding="utf-8")
    assert "paper/submission_minimal/" in readme_text
    assert "manuscript/" in readme_text


def test_create_submission_minimal_package_frontiers_family_profile_creates_journal_specific_assets(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    frontiers_root = tmp_path / "frontiers_resources"
    manuscript_template = frontiers_root / "Frontiers_Template.docx"
    supplementary_template = frontiers_root / "Supplementary_Material.docx"
    csl_path = frontiers_root / "frontiers.csl"

    write_docx(manuscript_template, "Frontiers manuscript template")
    write_docx(supplementary_template, "Frontiers supplementary template")
    csl_path.write_text(module.default_ama_csl_path().read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_TEMPLATE_DOCX", str(manuscript_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_SUPPLEMENTARY_TEMPLATE_DOCX", str(supplementary_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_CSL", str(csl_path))

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="frontiers_family_harvard",
    )

    submission_root = paper_root / "journal_submissions" / "frontiers_family_harvard"
    assert submission_root.exists()
    assert manifest["publication_profile"] == "frontiers_family_harvard"
    assert manifest["citation_style"] == "FrontiersHarvard"
    assert manifest["readme_path"] == "paper/journal_submissions/frontiers_family_harvard/README.md"
    assert (submission_root / "manuscript.docx").exists()
    assert (submission_root / "Supplementary_Material.docx").exists()
    assert (submission_root / "paper.pdf").exists()
    assert "paper/journal_submissions/frontiers_family_harvard/" in (
        submission_root / "README.md"
    ).read_text(encoding="utf-8")
    assert manifest["manuscript"]["docx_path"] == "paper/journal_submissions/frontiers_family_harvard/manuscript.docx"
    assert (
        manifest["supplementary_material"]["docx_path"]
        == "paper/journal_submissions/frontiers_family_harvard/Supplementary_Material.docx"
    )
    assert manifest["journal_target"]["journal_family"] == "Frontiers"
    assert manifest["journal_target"]["reference_style_family"] == "FrontiersHarvard"


def test_create_submission_minimal_package_rejects_legacy_frontiers_profile(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    frontiers_root = tmp_path / "frontiers_resources"
    manuscript_template = frontiers_root / "Frontiers_Template.docx"
    supplementary_template = frontiers_root / "Supplementary_Material.docx"
    csl_path = frontiers_root / "frontiers.csl"

    write_docx(manuscript_template, "Frontiers manuscript template")
    write_docx(supplementary_template, "Frontiers supplementary template")
    csl_path.write_text(module.default_ama_csl_path().read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_TEMPLATE_DOCX", str(manuscript_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_SUPPLEMENTARY_TEMPLATE_DOCX", str(supplementary_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_CSL", str(csl_path))

    try:
        module.create_submission_minimal_package(
            paper_root=paper_root,
            publication_profile="frontiers_in_physiology",
        )
    except ValueError as exc:
        assert "unsupported publication profile" in str(exc)
    else:
        raise AssertionError("legacy Frontiers profile should be rejected")


def test_create_submission_minimal_package_frontiers_family_profile_preserves_reference_doc_parts(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    frontiers_root = tmp_path / "frontiers_resources"
    manuscript_template = frontiers_root / "Frontiers_Template.docx"
    supplementary_template = frontiers_root / "Supplementary_Material.docx"
    csl_path = frontiers_root / "frontiers.csl"

    write_docx(manuscript_template, "Frontiers manuscript template")
    write_docx(supplementary_template, "Frontiers supplementary template")
    csl_path.write_text(module.default_ama_csl_path().read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_TEMPLATE_DOCX", str(manuscript_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_SUPPLEMENTARY_TEMPLATE_DOCX", str(supplementary_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_CSL", str(csl_path))

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="frontiers_family_harvard",
    )

    output_docx_path = paper_root / "journal_submissions" / "frontiers_family_harvard" / "manuscript.docx"
    with zipfile.ZipFile(output_docx_path) as archive:
        footer_names = [name for name in archive.namelist() if name.startswith("word/footer")]
        assert footer_names


def test_create_submission_minimal_package_frontiers_family_uses_figure_semantics_manifest_for_legends(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    frontiers_root = tmp_path / "frontiers_resources"
    manuscript_template = frontiers_root / "Frontiers_Template.docx"
    supplementary_template = frontiers_root / "Supplementary_Material.docx"
    csl_path = frontiers_root / "frontiers.csl"

    write_docx(manuscript_template, "Frontiers manuscript template")
    write_docx(supplementary_template, "Frontiers supplementary template")
    csl_path.write_text(module.default_ama_csl_path().read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_TEMPLATE_DOCX", str(manuscript_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_SUPPLEMENTARY_TEMPLATE_DOCX", str(supplementary_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_CSL", str(csl_path))

    dump_json(
        paper_root / "figure_semantics_manifest.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "story_role": "overall_performance_and_clinical_utility",
                    "research_question": "Does the clinically informed model improve utility beyond the reference model?",
                    "direct_message": "Calibration and clinical utility improved, whereas discrimination gains were modest.",
                    "clinical_implication": "Supports preoperative counseling and postoperative surveillance planning.",
                    "interpretation_boundary": "This figure does not establish a recommended intervention threshold.",
                    "panel_messages": [
                        {"panel_id": "A", "message": "Discrimination is only one component of the figure-level interpretation."}
                    ],
                    "legend_glossary": [
                        {
                            "term": "treat all",
                            "explanation": "Assumes every patient is managed as high risk at the chosen threshold."
                        },
                        {
                            "term": "treat none",
                            "explanation": "Assumes no patient is managed as high risk at the chosen threshold."
                        },
                    ],
                    "threshold_semantics": "Thresholds are illustrative operating points rather than recommended cut-offs.",
                    "stratification_basis": "Risk groups are display-oriented rather than prespecified clinical bins.",
                    "recommendation_boundary": "No formal threshold recommendation is made from this figure alone.",
                    "renderer_contract": {
                        "figure_semantics": "evidence",
                        "renderer_family": "python",
                        "selection_rationale": "The legend is derived from an audited evidence figure exported from the locked analysis stack.",
                        "fallback_on_failure": False,
                        "failure_action": "block_and_fix_environment"
                    }
                }
            ],
        },
    )

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="frontiers_family_harvard",
    )

    frontiers_markdown = (
        paper_root / "journal_submissions" / "frontiers_family_harvard" / "frontiers_manuscript.md"
    ).read_text(encoding="utf-8")
    assert "Calibration and clinical utility improved" in frontiers_markdown
    assert "treat all" in frontiers_markdown
    assert "Assumes every patient is managed as high risk" in frontiers_markdown
    assert "illustrative operating points rather than recommended cut-offs" in frontiers_markdown
    assert "No formal threshold recommendation is made from this figure alone" in frontiers_markdown


def test_create_submission_minimal_package_frontiers_family_syncs_into_study_family_package(
    tmp_path: Path,
    monkeypatch,
) -> None:
    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_paper_workspace(tmp_path)
    # keep this assertion impossible for current implementation so the new sync contract is explicit
    frontiers_root = tmp_path / "frontiers_resources"
    manuscript_template = frontiers_root / "Frontiers_Template.docx"
    supplementary_template = frontiers_root / "Supplementary_Material.docx"
    csl_path = frontiers_root / "frontiers.csl"

    write_docx(manuscript_template, "Frontiers manuscript template")
    write_docx(supplementary_template, "Frontiers supplementary template")
    csl_path.write_text(module.default_ama_csl_path().read_text(encoding="utf-8"), encoding="utf-8")

    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_TEMPLATE_DOCX", str(manuscript_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_SUPPLEMENTARY_TEMPLATE_DOCX", str(supplementary_template))
    monkeypatch.setenv("DEEPSCIENTIST_FRONTIERS_CSL", str(csl_path))

    called: dict[str, object] = {}

    def fake_can_sync(*, paper_root: Path) -> bool:
        return True

    def fake_sync(*, paper_root: Path, stage: str, publication_profile: str = "general_medical_journal") -> dict:
        called["paper_root"] = paper_root
        called["stage"] = stage
        called["publication_profile"] = publication_profile
        return {"stage": stage, "publication_profile": publication_profile}

    monkeypatch.setattr(module.study_delivery_sync, "can_sync_study_delivery", fake_can_sync)
    monkeypatch.setattr(module.study_delivery_sync, "sync_study_delivery", fake_sync)

    module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="frontiers_family_harvard",
    )

    assert called["paper_root"] == paper_root
    assert called["stage"] == "submission_minimal"
    assert called["publication_profile"] == "frontiers_family_harvard"


def test_create_submission_minimal_package_builds_submission_facing_docx_for_current_draft_shape(
    tmp_path: Path,
) -> None:
    from docx import Document

    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_current_draft_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    compiled_submission_markdown = submission_root / "manuscript_submission.md"
    assert compiled_submission_markdown.exists()

    submission_markdown = compiled_submission_markdown.read_text(encoding="utf-8")
    assert submission_markdown.startswith("---\n")
    assert 'title: "Current Draft Title"' in submission_markdown
    assert "bibliography: ../references.bib" in submission_markdown
    assert "\n# Abstract\n" in submission_markdown
    assert "\n# Conclusion\n" in submission_markdown
    assert not submission_markdown.startswith("# Draft")
    assert manifest["manuscript"]["source_markdown_path"] == "paper/submission_minimal/manuscript_submission.md"

    document = Document(submission_root / "manuscript.docx")
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    assert paragraphs[0] == "Current Draft Title"
    assert "Draft" not in paragraphs[:3]
    assert any("A primary source" in paragraph for paragraph in paragraphs)
    assert not any("ref1?" in paragraph for paragraph in paragraphs)


def test_create_submission_minimal_package_supports_manuscript_shaped_draft_without_front_matter(
    tmp_path: Path,
) -> None:
    from docx import Document

    module = importlib.import_module("med_autoscience.controllers.submission_minimal")
    paper_root = make_manuscript_shaped_draft_workspace(tmp_path)

    manifest = module.create_submission_minimal_package(
        paper_root=paper_root,
        publication_profile="general_medical_journal",
    )

    submission_root = paper_root / "submission_minimal"
    compiled_submission_markdown = submission_root / "manuscript_submission.md"
    assert compiled_submission_markdown.exists()

    submission_markdown = compiled_submission_markdown.read_text(encoding="utf-8")
    assert submission_markdown.startswith("---\n")
    assert 'title: "Manuscript-Shaped Draft Title"' in submission_markdown
    assert "bibliography: ../references.bib" in submission_markdown
    assert "bibliography: ../../references.bib" not in submission_markdown
    assert "title: \"Article Title\"" not in submission_markdown
    assert "\n# Methods\n\nStudy methods paragraph.\n" in submission_markdown
    assert "Draft abstract methods." in submission_markdown
    assert manifest["manuscript"]["source_markdown_path"] == "paper/submission_minimal/manuscript_submission.md"

    document = Document(submission_root / "manuscript.docx")
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    assert paragraphs[0] == "Manuscript-Shaped Draft Title"
    assert any("Study methods paragraph." in paragraph for paragraph in paragraphs)
    assert any("A primary source" in paragraph for paragraph in paragraphs)
