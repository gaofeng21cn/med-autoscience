from __future__ import annotations

import importlib
import io
import json
import os
from pathlib import Path
import shutil
import zipfile
import zlib

import pytest
from pypdf import PdfReader


_LIGHTWEIGHT_DOCX_BY_IMAGE_COUNT: dict[int, bytes] = {}
_LIGHTWEIGHT_PDF_BY_IMAGE_COUNT: dict[int, bytes] = {}


def dump_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def write_png(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(_png_bytes())


def write_open_authority_snapshots(study_root: Path) -> None:
    dump_json(
        study_root / "artifacts" / "truth" / "latest.json",
        {
            "surface": "study_truth_snapshot",
            "study_id": study_root.name,
            "truth_epoch": "truth-1",
            "canonical_next_action": "continue_bundle_stage",
            "allowed_controller_actions": [
                "direct_study_execution",
                "direct_paper_line_write",
                "direct_bundle_build",
                "direct_compiled_bundle_proofing",
            ],
            "blocking_reasons": [],
        },
    )
    dump_json(
        study_root / "artifacts" / "runtime" / "health" / "latest.json",
        {
            "surface": "runtime_health_snapshot",
            "study_id": study_root.name,
            "quest_id": study_root.name,
            "runtime_health_epoch": "runtime-1",
            "canonical_runtime_action": "continue_supervising_runtime",
            "attempt_state": "idle",
            "retry_budget_remaining": 1,
            "blocking_reasons": [],
        },
    )


def remove_authority_snapshots(study_root: Path) -> None:
    for path in (
        study_root / "artifacts" / "truth" / "latest.json",
        study_root / "artifacts" / "runtime" / "health" / "latest.json",
        study_root / "artifacts" / "publication_eval" / "latest.json",
    ):
        path.unlink(missing_ok=True)


def _png_bytes() -> bytes:
    raw_scanline = b"\x00\xff\xff\xff"
    compressed = zlib.compress(raw_scanline)

    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(chunk_type + data) & 0xFFFFFFFF
        return len(data).to_bytes(4, "big") + chunk_type + data + crc.to_bytes(4, "big")

    png_bytes = b"".join(
        [
            b"\x89PNG\r\n\x1a\n",
            chunk(b"IHDR", b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x02\x00\x00\x00"),
            chunk(b"IDAT", compressed),
            chunk(b"IEND", b""),
        ]
    )
    return png_bytes


def _markdown_image_count(path: Path) -> int:
    if not path.exists():
        return 0
    return sum(1 for line in path.read_text(encoding="utf-8").splitlines() if line.strip().startswith("!["))


def _lightweight_docx_bytes(image_count: int) -> bytes:
    resolved_image_count = max(0, int(image_count))
    cached = _LIGHTWEIGHT_DOCX_BY_IMAGE_COUNT.get(resolved_image_count)
    if cached is not None:
        return cached

    from docx import Document

    document = Document()
    document.add_paragraph("Lightweight submission export fixture")
    document.sections[0].footer.paragraphs[0].text = "Lightweight submission export fixture footer"
    png = _png_bytes()
    for _ in range(resolved_image_count):
        document.add_picture(io.BytesIO(png))
    output = io.BytesIO()
    document.save(output)
    payload = output.getvalue()
    _LIGHTWEIGHT_DOCX_BY_IMAGE_COUNT[resolved_image_count] = payload
    return payload


def _lightweight_pdf_bytes(image_count: int) -> bytes:
    resolved_image_count = max(0, int(image_count))
    cached = _LIGHTWEIGHT_PDF_BY_IMAGE_COUNT.get(resolved_image_count)
    if cached is not None:
        return cached

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    panel_count = max(1, resolved_image_count)
    figure, axes = plt.subplots(panel_count, 1, figsize=(1, panel_count))
    if panel_count == 1:
        axes = [axes]
    for index, axis in enumerate(axes):
        axis.axis("off")
        if index < resolved_image_count:
            axis.imshow([[(1.0, 1.0, 1.0)]])
    output = io.BytesIO()
    figure.savefig(output, format="pdf")
    plt.close(figure)
    payload = output.getvalue()
    _LIGHTWEIGHT_PDF_BY_IMAGE_COUNT[resolved_image_count] = payload
    return payload


@pytest.fixture
def real_submission_exports() -> None:
    pass


@pytest.fixture(autouse=True)
def lightweight_submission_exports(request: pytest.FixtureRequest, monkeypatch: pytest.MonkeyPatch) -> None:
    if "real_submission_exports" in request.fixturenames:
        return

    package_builder = importlib.import_module("med_autoscience.controllers.submission_minimal_parts.package_builder")
    submission_minimal = importlib.import_module("med_autoscience.controllers.submission_minimal")

    def export_docx(
        *,
        compiled_markdown_path: Path,
        paper_root: Path,
        output_docx_path: Path,
        csl_path: Path,
        reference_doc_path: Path | None = None,
    ) -> None:
        del paper_root, csl_path, reference_doc_path
        output_docx_path.parent.mkdir(parents=True, exist_ok=True)
        output_docx_path.write_bytes(_lightweight_docx_bytes(_markdown_image_count(compiled_markdown_path)))

    def export_pdf(
        *,
        compiled_markdown_path: Path,
        paper_root: Path,
        output_pdf_path: Path,
        csl_path: Path,
    ) -> None:
        del paper_root, csl_path
        output_pdf_path.parent.mkdir(parents=True, exist_ok=True)
        output_pdf_path.write_bytes(_lightweight_pdf_bytes(_markdown_image_count(compiled_markdown_path)))

    monkeypatch.setattr(package_builder, "export_docx", export_docx)
    monkeypatch.setattr(package_builder, "export_pdf", export_pdf)
    monkeypatch.setattr(submission_minimal, "export_docx", export_docx)
    monkeypatch.setattr(submission_minimal, "export_pdf", export_pdf)


def write_docx(path: Path, text: str) -> None:
    from docx import Document

    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph(text)
    document.sections[0].footer.paragraphs[0].text = f"{text} footer"
    document.save(path)


def make_paper_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "build" / "review_manuscript.md",
        """---
title: "Test Medical Manuscript"
bibliography: ../references.bib
link-citations: true
---

# Abstract

This is a manuscript citation [@ref1].

# Main Figures

## Figure 1. Main figure

Caption.

![](../figures/F1_main.png)

## Supplementary Figure S1. Supplementary figure

Caption.

![](../figures/FS1_supp.png)

# Main Tables

| Characteristic | Value |
| --- | --- |
| Age | 52 |
""",
    )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_png(paper_root / "figures" / "FS1_supp.png")
    write_text(paper_root / "figures" / "FS1_supp.pdf", "%PDF-1.4\n%supp figure\n")
    write_text(paper_root / "tables" / "T1_summary.csv", "Characteristic,Value\nAge,52\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")
    dump_json(
        paper_root / "evidence_ledger.json",
        {
            "schema_version": 1,
            "claims": [
                {
                    "claim_id": "C1",
                    "status": "supported",
                    "evidence": [{"source_paths": ["paper/figures/F1_main.png"]}],
                }
            ],
        },
    )

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "source_markdown": "paper/build/review_manuscript.md",
            "output_pdf": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "export_paths": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                },
                {
                    "figure_id": "FS1",
                    "paper_role": "supplementary",
                    "title": "Supplementary figure",
                    "export_paths": ["paper/figures/FS1_supp.pdf", "paper/figures/FS1_supp.png"],
                },
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "title": "Summary table",
                    "asset_paths": ["paper/tables/T1_summary.csv", "paper/tables/T1_summary.md"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "bundle_inputs": {
                "compiled_markdown_path": "paper/build/review_manuscript.md",
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
                {"path": "paper/figures/F1_main.pdf", "kind": "figure_export", "status": "present"},
                {"path": "paper/figures/FS1_supp.pdf", "kind": "supplementary_figure_export", "status": "present"},
                {"path": "paper/tables/T1_summary.csv", "kind": "table_asset", "status": "present"},
            ],
        },
    )

    write_open_authority_snapshots(workspace_root)
    return paper_root


def make_current_draft_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "draft.md",
        """# Draft

## Title

Current Draft Title

## Abstract

### Aims

Draft abstract aim with evidence [@ref1].

### Methods

Draft abstract methods.

### Results

Draft abstract results.

### Conclusions

Draft abstract conclusions.

## Introduction

Intro paragraph with citation [@ref1].

## Methods

Methods paragraph.

## Results

Results paragraph.

## Discussion

Discussion paragraph.

## Conclusion

Conclusion paragraph.
""",
        )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A. and Author, B.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "pdf_path": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "planned_exports": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "path": "paper/tables/T1_summary.md",
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/draft.md",
            "compile_report_path": "paper/build/compile_report.json",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
            ],
        },
    )
    write_open_authority_snapshots(workspace_root)
    return paper_root


def make_manuscript_shaped_draft_workspace(tmp_path: Path) -> Path:
    workspace_root = tmp_path / "workspace"
    paper_root = workspace_root / "paper"

    write_text(
        paper_root / "draft.md",
        """# Manuscript-Shaped Draft Title

## Abstract

### Background

Draft abstract background with evidence [@ref1].

### Methods

Draft abstract methods.

### Results

Draft abstract results.

### Conclusions

Draft abstract conclusions.

## Introduction

Intro paragraph with citation [@ref1].

## Materials and Methods

Study methods paragraph.

## Results

Results paragraph.

## Discussion

Discussion paragraph.
""",
    )
    write_text(
        paper_root / "references.bib",
        """@article{ref1,
  title={A primary source},
  author={Author, A. and Author, B.},
  journal={Journal},
  year={2024}
}
""",
    )
    write_png(paper_root / "figures" / "F1_main.png")
    write_text(paper_root / "figures" / "F1_main.pdf", "%PDF-1.4\n%main figure\n")
    write_text(paper_root / "tables" / "T1_summary.md", "| Characteristic | Value |\n| --- | --- |\n| Age | 52 |\n")
    write_text(paper_root / "paper.pdf", "%PDF-1.4\n%paper bundle\n")

    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "pdf_path": "paper/paper.pdf",
        },
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Main figure",
                    "planned_exports": ["paper/figures/F1_main.pdf", "paper/figures/F1_main.png"],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "path": "paper/tables/T1_summary.md",
                }
            ],
        },
    )
    dump_json(
        paper_root / "paper_bundle_manifest.json",
        {
            "schema_version": 1,
            "draft_path": "paper/draft.md",
            "compile_report_path": "paper/build/compile_report.json",
            "bundle_inputs": {
                "compile_report_path": "paper/build/compile_report.json",
                "figure_catalog_path": "paper/figures/figure_catalog.json",
                "table_catalog_path": "paper/tables/table_catalog.json",
            },
            "included_assets": [
                {"path": "paper/paper.pdf", "kind": "compiled_pdf", "status": "present"},
            ],
        },
    )
    write_open_authority_snapshots(workspace_root)
    return paper_root


def make_materialized_submission_source_workspace(tmp_path: Path) -> Path:
    paper_root = make_current_draft_workspace(tmp_path)
    write_png(paper_root / "submission_minimal" / "figures" / "F1.png")

    write_text(
        paper_root / "submission_minimal" / "manuscript_source.md",
        """# Materialized Submission Title

## Abstract

Materialized abstract with evidence [@ref1].

## Introduction

Materialized introduction with evidence [@ref1].

## Methods

Materialized methods paragraph.

## Results

Materialized results paragraph.

## Discussion

Materialized discussion paragraph.

## Conclusion

Materialized conclusion paragraph.

## Main-text figures

### F1. Main figure

![F1](figures/F1.png)

Materialized figure caption.
""",
    )
    dump_json(
        paper_root / "build" / "compile_report.json",
        {
            "source_markdown_path": "paper/submission_minimal/manuscript_source.md",
            "pdf_path": "paper/paper.pdf",
        },
    )
    return paper_root


def make_authoritative_worktree_source_workspace(tmp_path: Path) -> Path:
    seed_paper_root = make_paper_workspace(tmp_path / "seed")
    paper_root = tmp_path / "quest" / ".ds" / "worktrees" / "paper-run-123" / "paper"
    shutil.copytree(seed_paper_root, paper_root, dirs_exist_ok=True)
    write_open_authority_snapshots(paper_root.parent)

    write_png(paper_root / "figures" / "generated" / "F1.png")
    write_text(paper_root / "figures" / "generated" / "F1.pdf", "%PDF-1.4\n%authoritative figure\n")
    write_text(paper_root / "tables" / "generated" / "T1.csv", "Characteristic,Value\nAge,99\n")
    write_text(
        paper_root / "tables" / "generated" / "T1.md",
        "| Characteristic | Value |\n| --- | --- |\n| Age | 99 |\n",
    )
    dump_json(
        paper_root / "figures" / "figure_catalog.json",
        {
            "schema_version": 1,
            "figures": [
                {
                    "figure_id": "F1",
                    "paper_role": "main_text",
                    "title": "Authoritative figure",
                    "export_paths": [
                        "../../../paper/figures/generated/F1.pdf",
                        "../../../paper/figures/generated/F1.png",
                    ],
                    "planned_exports": [
                        "paper/figures/generated/F1.pdf",
                        "paper/figures/generated/F1.png",
                    ],
                }
            ],
        },
    )
    dump_json(
        paper_root / "tables" / "table_catalog.json",
        {
            "schema_version": 1,
            "tables": [
                {
                    "table_id": "T1",
                    "paper_role": "main_text",
                    "title": "Authoritative table",
                    "asset_paths": [
                        "../../../paper/tables/generated/T1.csv",
                        "../../../paper/tables/generated/T1.md",
                    ],
                }
            ],
        },
    )
    return paper_root


def make_stage_native_current_body_workspace(tmp_path: Path) -> Path:
    seed_paper_root = make_current_draft_workspace(tmp_path / "seed")
    workspace_root = tmp_path / "workspace"
    study_root = workspace_root / "studies" / "003-dpcc"
    paper_root = (
        study_root
        / "artifacts"
        / "stage_outputs"
        / "_body_authority"
        / "paper_authority_cutover"
        / "current_body"
        / "paper"
    )
    shutil.copytree(seed_paper_root, paper_root, dirs_exist_ok=True)
    dump_json(
        paper_root / "evidence_ledger.json",
        {
            "schema_version": 1,
            "claims": [
                {
                    "claim_id": "C1",
                    "status": "supported",
                    "evidence_refs": ["paper/draft.md"],
                }
            ],
        },
    )
    (study_root / "study.yaml").write_text("study_id: 003-dpcc\n", encoding="utf-8")
    (study_root / "runtime_binding.yaml").write_text(
        "\n".join(
            [
                "schema_version: 1",
                "study_id: 003-dpcc",
                "quest_id: 003-dpcc-managed",
                "",
            ]
        ),
        encoding="utf-8",
    )
    quest_root = workspace_root / "runtime" / "quests" / "003-dpcc-managed"
    quest_root.mkdir(parents=True, exist_ok=True)
    (quest_root / "quest.yaml").write_text(
        "quest_id: 003-dpcc-managed\n"
        "runtime_reentry_gate:\n"
        "  study_id: 003-dpcc\n",
        encoding="utf-8",
    )
    write_open_authority_snapshots(study_root)
    return paper_root
